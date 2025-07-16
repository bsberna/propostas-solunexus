# streamlit app finalizado (adicione o conteúdo completo aqui)
import streamlit as st
from docx import Document
import datetime
import pypandoc
import os
import pandas as pd
import json

# === Arquivo de histórico ===
historico_path = "propostas_emitidas.json"
usuarios_path = "usuarios.json"
emails_path = "emails.json"

# === Carregar dados persistentes ===
def carregar_dados(path, default):
    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    else:
        with open(path, "w") as f:
            json.dump(default, f)
        return default

usuarios = carregar_dados(usuarios_path, {"vendedor1": "senha123"})
emails = carregar_dados(emails_path, {"vendedor1": "vendedor1@email.com"})
historico = carregar_dados(historico_path, [])

# === Salvar dados atualizados ===
def salvar_dados(path, dados):
    with open(path, "w") as f:
        json.dump(dados, f, indent=2)

# === Substituir campos no documento ===
def substituir_campos(doc, campos):
    for p in doc.paragraphs:
        for k, v in campos.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for k, v in campos.items():
                    if k in c.text:
                        c.text = c.text.replace(k, v)

# === Gerar PDF do histórico ===
def gerar_pdf_historico(df):
    html = f"<h2>Histórico de Propostas</h2>{df.to_html(index=False)}"
    with open("historico.html", "w", encoding="utf-8") as f:
        f.write(html)
    output = "historico_propostas.pdf"
    pypandoc.convert_file("historico.html", 'pdf', outputfile=output)
    return output

# === Registrar proposta emitida ===
def registrar_proposta(info):
    historico.append(info)
    salvar_dados(historico_path, historico)

# === Tela de histórico ===
def visualizar_historico():
    st.subheader("📚 Histórico de Propostas Emitidas")
    if not historico:
        st.info("Nenhuma proposta registrada ainda.")
        return
    df = pd.DataFrame(historico)
    busca = st.text_input("🔍 Buscar por cliente ou código")
    vendedor = st.selectbox("👤 Filtrar por vendedor", ["Todos"] + sorted(df["usuario"].unique()))
    if busca:
        df = df[df.apply(lambda r: busca.lower() in str(r.values).lower(), axis=1)]
    if vendedor != "Todos":
        df = df[df["usuario"] == vendedor]
    st.dataframe(df)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("⬇️ CSV", df.to_csv(index=False), "historico.csv")
    with col2:
        st.download_button("⬇️ JSON", df.to_json(orient="records", indent=2), "historico.json")
    with col3:
        pdf_path = gerar_pdf_historico(df)
        with open(pdf_path, "rb") as f:
            st.download_button("📄 PDF", f, file_name="historico_propostas.pdf")

    st.markdown("### 📥 Propostas Individuais")
    for _, row in df.iterrows():
        if os.path.exists(row["pdf"]):
            with open(row["pdf"], "rb") as f:
                st.download_button(f"{row['codigo']} - {row['cliente']}", f, file_name=os.path.basename(row["pdf"]))

# === Geração da proposta ===
def gerar_proposta(campos):
    modelo = "Proposta Comercial xxx.x.xxxx.docx" if campos["TIPO"] == "Comercial" else "Proposta Técnica xxx.x.xxxx.docx"
    nome_base = f"Proposta_{campos['xxx.x.xxxx']}_{campos['Cliente']}"
    saida_docx = nome_base + ".docx"
    saida_pdf = nome_base + ".pdf"

    doc = Document(modelo)
    substituir_campos(doc, campos)
    doc.save(saida_docx)
    pypandoc.convert_file(saida_docx, "pdf", outputfile=saida_pdf)

    with open(saida_pdf, "rb") as f:
        st.download_button("📄 Baixar Proposta PDF", f, file_name=saida_pdf)

    registrar_proposta({
        "codigo": campos["xxx.x.xxxx"],
        "cliente": campos["Cliente"],
        "tipo": campos["TIPO"],
        "usuario": st.session_state.get("usuario", ""),
        "data": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "pdf": saida_pdf
    })

# === Formulário da proposta ===
def tela_proposta():
    st.subheader("✏️ Preencher Proposta")
    campos = {
        "TIPO": st.selectbox("Tipo de proposta", ["Comercial", "Técnica"]),
        "xxx.x.xxxx": st.text_input("Código da proposta (ex: 019.1.2025)"),
        "Cliente": st.text_input("Nome do cliente"),
        "Serviço": st.text_input("Serviço contratado"),
        "Área": st.text_input("Área em hectares")
    }
    if st.button("✅ Gerar Proposta"):
        gerar_proposta(campos)

# === Autenticação e cadastro ===
def autenticar():
    aba = st.radio("Acesso", ["Entrar", "Esqueci minha senha", "Cadastrar novo usuário"])
    if aba == "Entrar":
        user = st.text_input("Usuário")
        senha = st.text_input("Senha", type="password")
        if st.button("Entrar"):
            if usuarios.get(user) == senha:
                st.session_state["logado"] = True
                st.session_state["usuario"] = user
            else:
                st.error("Usuário ou senha incorretos.")
    elif aba == "Esqueci minha senha":
        email = st.text_input("Digite seu e-mail")
        if st.button("Enviar instruções"):
            if email in emails.values():
                st.success(f"Instruções enviadas para {email} (simulado).")
            else:
                st.error("E-mail não encontrado.")
    elif aba == "Cadastrar novo usuário":
        novo = st.text_input("Novo usuário")
        senha = st.text_input("Senha", type="password")
        email = st.text_input("E-mail")
        if st.button("Cadastrar"):
            if novo in usuarios:
                st.error("Usuário já existe.")
            else:
                usuarios[novo] = senha
                emails[novo] = email
                salvar_dados(usuarios_path, usuarios)
                salvar_dados(emails_path, emails)
                st.success(f"{novo} cadastrado com sucesso!")

# === Início do app ===
st.set_page_config(page_title="Propostas Solunexus", layout="wide")

if "logado" not in st.session_state:
    st.session_state["logado"] = False

if not st.session_state["logado"]:
    autenticar()
else:
    st.sidebar.success(f"Logado como {st.session_state['usuario']}")
    pagina = st.sidebar.selectbox("Menu", ["Gerar Proposta", "Histórico", "Sair"])
    if pagina == "Gerar Proposta":
        tela_proposta()
    elif pagina == "Histórico":
        visualizar_historico()
    elif pagina == "Sair":
        st.session_state["logado"] = False
        st.experimental_rerun()
